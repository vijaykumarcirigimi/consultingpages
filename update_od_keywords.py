"""
Script to update Organizational Consulting-Final.docx with SEO keywords.
Every keyword phrase inserted is bolded.
"""

from docx import Document
from docx.shared import RGBColor
import re
import copy

INPUT_FILE = 'content/Final Pages/Organizational Consulting-Final.docx'
OUTPUT_FILE = 'content/Final Pages/Organizational Consulting-Final.docx'

doc = Document(INPUT_FILE)

# ─── Helper: replace text in a paragraph, bolding keyword phrases ───
def replace_in_paragraph_with_bold(paragraph, old_text, new_text, keywords_to_bold):
    """Replace old_text with new_text in paragraph, bolding any keyword phrases found in new_text."""
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Capture formatting from first run
    font_name = font_size = font_color = font_bold = None
    if paragraph.runs:
        r = paragraph.runs[0]
        font_name = r.font.name
        font_size = r.font.size
        font_color = r.font.color.rgb if r.font.color and r.font.color.rgb else None
        font_bold = r.bold

    new_full_text = full_text.replace(old_text, new_text, 1)

    # Clear existing runs
    p_el = paragraph._element
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for r in p_el.findall(f'{ns}r'):
        p_el.remove(r)

    # Build regex pattern from all keywords to bold (longest first to avoid partial matches)
    if isinstance(keywords_to_bold, str):
        keywords_to_bold = [keywords_to_bold]
    sorted_kws = sorted(keywords_to_bold, key=len, reverse=True)
    pattern = '(' + '|'.join(re.escape(kw) for kw in sorted_kws) + ')'

    parts = re.split(pattern, new_full_text, flags=re.IGNORECASE)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_color:
            run.font.color.rgb = font_color
        # Check if this part matches any keyword
        is_keyword = any(part.lower() == kw.lower() for kw in sorted_kws)
        if is_keyword:
            run.bold = True
        elif font_bold:
            run.bold = font_bold
    return True


def replace_in_cell_with_bold(cell, old_text, new_text, keywords_to_bold):
    """Apply replacement across all paragraphs in a table cell."""
    changed = False
    for para in cell.paragraphs:
        if replace_in_paragraph_with_bold(para, old_text, new_text, keywords_to_bold):
            changed = True
    return changed


def replace_in_table_cell0(table_index, old_text, new_text, keywords_to_bold):
    """Replace in the first cell (row 0, col 0) of a table."""
    if table_index >= len(doc.tables):
        print(f"  WARNING: Table {table_index} does not exist")
        return False
    cell = doc.tables[table_index].rows[0].cells[0]
    result = replace_in_cell_with_bold(cell, old_text, new_text, keywords_to_bold)
    if result:
        print(f"  T{table_index}: replaced '{old_text[:60]}...' -> keyword-enriched")
    return result


def replace_in_table_cell(table_index, row, col, old_text, new_text, keywords_to_bold):
    """Replace in a specific cell of a table."""
    if table_index >= len(doc.tables):
        print(f"  WARNING: Table {table_index} does not exist")
        return False
    t = doc.tables[table_index]
    if row >= len(t.rows) or col >= len(t.columns):
        print(f"  WARNING: T{table_index} R{row}C{col} does not exist")
        return False
    cell = t.rows[row].cells[col]
    result = replace_in_cell_with_bold(cell, old_text, new_text, keywords_to_bold)
    if result:
        print(f"  T{table_index} R{row}C{col}: replaced '{old_text[:50]}...'")
    return result


# ─── All keywords we want to bold when they appear ───
ALL_KEYWORDS = [
    "organizational development consulting services",
    "organizational development consulting solutions",
    "organizational development consulting company",
    "organizational development consulting",
    "organizational development consultant",
    "organizational development consulting provider",
    "organizational development consulting firm",
    "organizational development consulting experts",
    "organizational development consulting outsourcing",
    "organizational development consultation",
    "organizational development consulting agency",
    "organizational development consulting vendor",
    "organizational development consulting partner",
    "hire organizational development consultant",
    "best organizational development consulting company",
    "top organizational development consulting providers",
    "organizational development consulting firms",
    "organizational development consulting service provider",
    "affordable organizational development consulting services",
    "professional organizational development consulting services",
    "certified organizational development consultants",
    "experienced organizational development consulting company",
    "corporate organizational development consulting services",
    "enterprise organizational development consulting solutions",
    "organizational development consulting for enterprises",
    "employee organizational development consulting provider",
    "workforce organizational development consulting company",
    "custom organizational development consulting services",
    "end-to-end organizational development consulting solutions",
    "managed organizational development consulting services",
    "outsource organizational development consulting",
    "organizational development and culture transformation consulting",
    "strategic organizational development consulting for large enterprises",
]

changes = 0

# ════════════════════════════════════════════════════════════
# T1: SEO Metadata
# ════════════════════════════════════════════════════════════
print("\n=== T1: SEO Metadata ===")

# Meta Title — already has "Organizational Development Consulting | Edstellar" — add "Services"
replace_in_table_cell(1, 1, 1,
    "Organizational Development Consulting | Edstellar",
    "Organizational Development Consulting Services | Edstellar",
    ["organizational development consulting services"])
changes += 1

# Meta Description — expand "OD consultants" to full keyword
replace_in_table_cell(1, 2, 1,
    "Partner with Edstellar's expert OD consultants to align people, culture & performance. 2,000+ instructor-led programs across 100+ locations.",
    "Partner with Edstellar's expert organizational development consultants to align people, culture & performance. Professional organizational development consulting services — 2,000+ instructor-led programs across 100+ locations.",
    ["organizational development consultants", "professional organizational development consulting services"])
changes += 1

# ════════════════════════════════════════════════════════════
# T2: Page Structure — weave keywords into section names/rationale
# ════════════════════════════════════════════════════════════
print("\n=== T2: Page Structure ===")

replace_in_table_cell(2, 2, 1,
    "What Is OD Consulting? (Definitional)",
    "What Is Organizational Development Consulting? (Definitional)",
    ["organizational development consulting"])
changes += 1

replace_in_table_cell(2, 6, 1,
    "OD Consulting Services (8 tiles)",
    "Organizational Development Consulting Services (8 tiles)",
    ["organizational development consulting services"])
changes += 1

replace_in_table_cell(2, 12, 1,
    "Our Approach to OD Consulting",
    "Our Approach to Organizational Development Consulting",
    ["organizational development consulting"])
changes += 1

replace_in_table_cell(2, 14, 1,
    "Why Choose Edstellar?",
    "Why Choose Edstellar for Organizational Development Consulting?",
    ["organizational development consulting"])
changes += 1

replace_in_table_cell(2, 19, 1,
    "Gated Asset — OD Readiness Checklist",
    "Gated Asset — Organizational Development Readiness Checklist",
    ["organizational development"])
changes += 1

# ════════════════════════════════════════════════════════════
# T5: Hero Section
# ════════════════════════════════════════════════════════════
print("\n=== T5: Hero Section ===")

# Expand value prop - replace "OD consultants" with full keyword and add more keywords
replace_in_table_cell0(5,
    "Partner with Edstellar's expert OD consultants to align people, culture, and performance. Evidence-backed organizational development consulting - driving culture, capability, and performance transformation through 2,000+ instructor-led programs, 5,000+ certified trainers, delivered across 100+ locations worldwide.",
    "Partner with Edstellar — a trusted organizational development consulting company — to align people, culture, and performance. Our certified organizational development consultants deliver evidence-backed organizational development consulting services, driving culture, capability, and performance transformation through 2,000+ instructor-led programs, 5,000+ certified trainers, delivered across 100+ locations worldwide.",
    ["organizational development consulting company", "certified organizational development consultants", "organizational development consulting services"])
changes += 1

# Badge - expand "OD Consulting"
replace_in_table_cell0(5,
    "OD Consulting  —  Enterprise Practice",
    "Organizational Development Consulting  —  Enterprise Practice",
    ["organizational development consulting"])
changes += 1

# ════════════════════════════════════════════════════════════
# T7: What Is Organizational Development Consulting?
# ════════════════════════════════════════════════════════════
print("\n=== T7: What Is OD Consulting? ===")

replace_in_table_cell0(7,
    "Organizational development consulting is a planned, systematic approach to improving organizational effectiveness through people, processes, and structure. Rooted in research-backed OD methodologies — including McKinsey 7S, Burke-Litwin, and systems thinking — OD consulting identifies root-cause barriers to performance and designs evidence-based interventions that drive lasting change.",
    "Organizational development consulting is a planned, systematic approach to improving organizational effectiveness through people, processes, and structure. As a professional organizational development consulting firm, Edstellar roots every engagement in research-backed OD methodologies — including McKinsey 7S, Burke-Litwin, and systems thinking. Our organizational development consulting services identify root-cause barriers to performance and design evidence-based interventions that drive lasting change.",
    ["organizational development consulting", "professional organizational development consulting firm", "organizational development consulting services"])
changes += 1

replace_in_table_cell0(7,
    "At Edstellar, our OD consultants partner with leadership teams to align organizational strategy with operational reality — building the structure, culture, and capability your business needs to scale, adapt, and outperform.",
    "At Edstellar, our experienced organizational development consultants partner with leadership teams to align organizational strategy with operational reality — delivering end-to-end organizational development consulting solutions that build the structure, culture, and capability your business needs to scale, adapt, and outperform.",
    ["experienced organizational development consultants", "end-to-end organizational development consulting solutions", "organizational development consulting"])
changes += 1

# ════════════════════════════════════════════════════════════
# T9: Why Work with Our OD Consultants?
# ════════════════════════════════════════════════════════════
print("\n=== T9: Why Work with Our OD Consultants? ===")

replace_in_table_cell0(9,
    "Organizational development is complex, context-dependent, and high-stakes. The right OD consultant brings three things your internal team often cannot — objectivity, a proven methodology, and cross-industry pattern recognition. Here's why enterprises choose Edstellar's OD consultants:",
    "Organizational development consulting is complex, context-dependent, and high-stakes. The right organizational development consultant brings three things your internal team often cannot — objectivity, a proven methodology, and cross-industry pattern recognition. Here's why enterprises hire organizational development consultants from Edstellar — a top organizational development consulting provider:",
    ["organizational development consulting", "organizational development consultant", "hire organizational development consultants", "organizational development consulting provider"])
changes += 1

# ════════════════════════════════════════════════════════════
# T11: Trust / Proof Bar
# ════════════════════════════════════════════════════════════
print("\n=== T11: Trust / Proof Bar ===")
# T11 is mostly a SEO annotation. Let's check what's there
# It mainly has: "research-backed organizational development" — let's enhance the SEO note area
# This table has limited editable body text, mostly SEO annotations

# ════════════════════════════════════════════════════════════
# T13: Three Pillars
# ════════════════════════════════════════════════════════════
print("\n=== T13: Three Pillars ===")

replace_in_table_cell0(13,
    "Effective organizational transformation happens through deliberate improvement efforts and strategic reinforcement of what makes your organization successful. At Edstellar, we collaborate with leadership teams to align structure, culture, and talent with your evolving strategic direction.",
    "Effective organizational transformation happens through deliberate improvement efforts and strategic reinforcement of what makes your organization successful. As a leading organizational development consulting partner, Edstellar collaborates with leadership teams to align structure, culture, and talent with your evolving strategic direction through custom organizational development consulting services.",
    ["organizational development consulting partner", "custom organizational development consulting services"])
changes += 1

# ════════════════════════════════════════════════════════════
# T15: 8 OD Consulting Services tiles
# ════════════════════════════════════════════════════════════
print("\n=== T15: OD Consulting Services ===")

replace_in_table_cell0(15,
    "Edstellar's OD consulting services span the full spectrum of organizational transformation — from structural design and culture change through to leadership capability and workforce development. Each service is grounded in diagnostic evidence, not assumptions.",
    "Edstellar's organizational development consulting services span the full spectrum of organizational transformation — from structural design and culture change through to leadership capability and workforce development. As a best organizational development consulting company, each service is grounded in diagnostic evidence, not assumptions.",
    ["organizational development consulting services", "best organizational development consulting company"])
changes += 1

# ════════════════════════════════════════════════════════════
# T17: OD Diagnostic Study
# ════════════════════════════════════════════════════════════
print("\n=== T17: OD Diagnostic Study ===")

replace_in_table_cell0(17,
    "Before you redesign, restructure, or retrain — you need to know exactly where the structural and capability fractures are. Edstellar's OD Diagnostic Study is a structured, data-backed assessment of your organization's structural effectiveness, capability gaps, cultural health, and operational alignment.",
    "Before you redesign, restructure, or retrain — you need to know exactly where the structural and capability fractures are. Edstellar's organizational development consultation begins with an OD Diagnostic Study — a structured, data-backed assessment of your organization's structural effectiveness, capability gaps, cultural health, and operational alignment. This is the foundation of our managed organizational development consulting services.",
    ["organizational development consultation", "managed organizational development consulting services"])
changes += 1

# ════════════════════════════════════════════════════════════
# T19: Org Design & Operating Model
# ════════════════════════════════════════════════════════════
print("\n=== T19: Org Design & Operating Model ===")

replace_in_table_cell0(19,
    "Your operating model should accelerate execution — not constrain it. Edstellar designs organizational structures and governance frameworks that translate strategic intent into operational reality. We rethink how work flows, where decisions are made, and how accountability is distributed.",
    "Your operating model should accelerate execution — not constrain it. As a leading organizational development consulting firm, Edstellar designs organizational structures and governance frameworks that translate strategic intent into operational reality. Our enterprise organizational development consulting solutions rethink how work flows, where decisions are made, and how accountability is distributed.",
    ["organizational development consulting firm", "enterprise organizational development consulting solutions"])
changes += 1

# ════════════════════════════════════════════════════════════
# T21: Role Architecture
# ════════════════════════════════════════════════════════════
print("\n=== T21: Role Architecture ===")

replace_in_table_cell0(21,
    "When roles are unclear, accountability fragments. Edstellar goes beyond job descriptions — we conduct rigorous Job-Task Analysis (JTA), assess role criticality, map competencies, and redesign work to eliminate redundancy while maximizing value creation.",
    "When roles are unclear, accountability fragments. As an experienced organizational development consulting company, Edstellar goes beyond job descriptions — our organizational development consultants conduct rigorous Job-Task Analysis (JTA), assess role criticality, map competencies, and redesign work to eliminate redundancy while maximizing value creation.",
    ["experienced organizational development consulting company", "organizational development consultants"])
changes += 1

# ════════════════════════════════════════════════════════════
# T23: Strategic Planning
# ════════════════════════════════════════════════════════════
print("\n=== T23: Strategic Planning ===")

replace_in_table_cell0(23,
    "Strategy fails when it stays in the boardroom. Edstellar creates a structured cascade from enterprise strategy through business unit, team, and individual objectives — ensuring every employee understands how their work connects to business outcomes.",
    "Strategy fails when it stays in the boardroom. Edstellar's strategic organizational development consulting for large enterprises creates a structured cascade from enterprise strategy through business unit, team, and individual objectives — ensuring every employee understands how their work connects to business outcomes.",
    ["strategic organizational development consulting for large enterprises"])
changes += 1

# ════════════════════════════════════════════════════════════
# T25: Skills Maturity Assessment
# ════════════════════════════════════════════════════════════
print("\n=== T25: Skills Maturity Assessment ===")

replace_in_table_cell0(25,
    "Skills Maturity measures how effectively your organization identifies, develops, deploys, and evolves workforce capabilities. Edstellar's five-level maturity model benchmarks where you are today — and builds a data-driven roadmap to skills-led competitive advantage.",
    "Skills Maturity measures how effectively your organization identifies, develops, deploys, and evolves workforce capabilities. As a workforce organizational development consulting company, Edstellar's five-level maturity model benchmarks where you are today — and builds a data-driven roadmap to skills-led competitive advantage through organizational development consulting solutions.",
    ["workforce organizational development consulting company", "organizational development consulting solutions"])
changes += 1

# ════════════════════════════════════════════════════════════
# T27: Our Approach to OD Consulting
# ════════════════════════════════════════════════════════════
print("\n=== T27: Our Approach ===")

replace_in_table_cell0(27,
    "Organizational transformation is complex and requires a structured approach that addresses every dimension of change. Edstellar's methodology combines rigorous diagnostic analysis with collaborative solution design — grounded in research-backed organizational development methodologies.",
    "Organizational transformation is complex and requires a structured approach that addresses every dimension of change. Edstellar's organizational development consulting methodology combines rigorous diagnostic analysis with collaborative solution design — grounded in research-backed organizational development methodologies. As a top organizational development consulting provider, we deliver end-to-end organizational development consulting solutions.",
    ["organizational development consulting", "top organizational development consulting provider", "end-to-end organizational development consulting solutions"])
changes += 1

# ════════════════════════════════════════════════════════════
# T29: Industries We Serve
# ════════════════════════════════════════════════════════════
print("\n=== T29: Industries We Serve ===")

replace_in_table_cell0(29,
    "Edstellar delivers organizational development consulting across diverse sectors — from regulated industries to high-growth technology firms and public sector organizations. Our APAC, Middle East, and India-specific OD expertise fills a gap all 7 top competitors leave open.",
    "Edstellar delivers organizational development consulting for enterprises across diverse sectors — from regulated industries to high-growth technology firms and public sector organizations. As a corporate organizational development consulting services provider, our APAC, Middle East, and India-specific organizational development consulting expertise fills a gap all 7 top competitors leave open.",
    ["organizational development consulting for enterprises", "corporate organizational development consulting services", "organizational development consulting"])
changes += 1

# ════════════════════════════════════════════════════════════
# T31: Why Choose Edstellar
# ════════════════════════════════════════════════════════════
print("\n=== T31: Why Choose Edstellar ===")

replace_in_table_cell0(31,
    "Edstellar is not a generalist consulting firm. We are an organizational development consulting partner with the scale, methodology, and sector depth to deliver transformation that sticks. Here's what sets us apart from every other OD consulting firm in the market:",
    "Edstellar is not a generalist consulting firm. We are an organizational development consulting partner with the scale, methodology, and sector depth to deliver transformation that sticks. Here's what sets us apart from every other organizational development consulting firm in the market — making Edstellar the best organizational development consulting company for enterprise transformation:",
    ["organizational development consulting partner", "organizational development consulting firm", "best organizational development consulting company"])
changes += 1

# ════════════════════════════════════════════════════════════
# T33: Case Studies
# ════════════════════════════════════════════════════════════
print("\n=== T33: Case Studies ===")

replace_in_table_cell0(33,
    "Our organizational development expertise delivers measurable impact across diverse sectors.",
    "Our organizational development consulting expertise delivers measurable impact across diverse sectors. Each case study showcases Edstellar's professional organizational development consulting services in action.",
    ["organizational development consulting", "professional organizational development consulting services"])
changes += 1

# ════════════════════════════════════════════════════════════
# T35: Testimonials
# ════════════════════════════════════════════════════════════
print("\n=== T35: Testimonials ===")

replace_in_table_cell0(35,
    "We don't just consult — we co-create. Read how organizations experienced tangible shifts with Edstellar's organizational development consulting.",
    "We don't just consult — we co-create. Read how organizations experienced tangible shifts with Edstellar's organizational development consulting services. As certified organizational development consultants, we deliver measurable transformation.",
    ["organizational development consulting services", "certified organizational development consultants"])
changes += 1

# ════════════════════════════════════════════════════════════
# T37: Logo Wall
# ════════════════════════════════════════════════════════════
print("\n=== T37: Logo Wall ===")

replace_in_table_cell0(37,
    "Leading organizations across industries trust Edstellar for organizational development consulting. Our track record includes Fortune 500 companies, innovative startups, and respected public sector organizations.",
    "Leading organizations across industries trust Edstellar as their organizational development consulting vendor. Our track record as an organizational development consulting agency includes Fortune 500 companies, innovative startups, and respected public sector organizations.",
    ["organizational development consulting vendor", "organizational development consulting agency"])
changes += 1

# ════════════════════════════════════════════════════════════
# T39: FAQ
# ════════════════════════════════════════════════════════════
print("\n=== T39: FAQ ===")

# Q1 answer - already has "organizational development consulting" - add more keywords
replace_in_table_cell0(39,
    "Organizational development consulting is a systematic, evidence-based process for improving organizational effectiveness through planned interventions in structure, culture, capability, and process. OD consultants diagnose root causes — not surface symptoms — and design tailored solutions that produce measurable business outcomes.",
    "Organizational development consulting is a systematic, evidence-based process for improving organizational effectiveness through planned interventions in structure, culture, capability, and process. Organizational development consultants diagnose root causes — not surface symptoms — and design tailored organizational development consulting solutions that produce measurable business outcomes.",
    ["organizational development consulting", "organizational development consultants", "organizational development consulting solutions"])
changes += 1

# Q2 - expand "OD consultants"
replace_in_table_cell0(39,
    "Edstellar's OD consultants apply validated frameworks including McKinsey 7S, the Burke-Litwin Causal Model, Kotter's 8-Step Change Model, and Organizational Network Analysis (ONA). Every methodology is selected based on your organization's specific challenge — not applied universally.",
    "Edstellar's organizational development consultants apply validated frameworks including McKinsey 7S, the Burke-Litwin Causal Model, Kotter's 8-Step Change Model, and Organizational Network Analysis (ONA). As an organizational development consulting experts team, every methodology is selected based on your organization's specific challenge — not applied universally.",
    ["organizational development consultants", "organizational development consulting experts"])
changes += 1

# Q4 - expand "OD consulting"
replace_in_table_cell0(39,
    "Edstellar's OD consulting covers: Change Management, Leadership Development, Organizational Design & Structure, Cultural Transformation, Strategic Planning, Performance Management, Succession Planning, Organizational Health Assessment, Role Architecture, Goal Alignment, and Skills Maturity — all backed by instructor-led program delivery across 100+ locations.",
    "Edstellar's organizational development consulting services cover: Change Management, Leadership Development, Organizational Design & Structure, Organizational Development and Culture Transformation Consulting, Strategic Planning, Performance Management, Succession Planning, Organizational Health Assessment, Role Architecture, Goal Alignment, and Skills Maturity — all backed by instructor-led program delivery across 100+ locations.",
    ["organizational development consulting services", "organizational development and culture transformation consulting"])
changes += 1

# Q5 - add keyword
replace_in_table_cell0(39,
    "Every Edstellar engagement defines KPIs at the outset — retention rates, productivity metrics, cultural health scores, leadership pipeline readiness, and goal alignment indices. We use a four-phase model (Assess → Diagnose → Design → Implement & Measure) to track and calibrate results throughout the engagement.",
    "Every Edstellar organizational development consulting engagement defines KPIs at the outset — retention rates, productivity metrics, cultural health scores, leadership pipeline readiness, and goal alignment indices. We use a four-phase model (Assess → Diagnose → Design → Implement & Measure) to track and calibrate results, making Edstellar an affordable organizational development consulting services provider that delivers measurable ROI.",
    ["organizational development consulting", "affordable organizational development consulting services"])
changes += 1

# Q6 - expand keywords
replace_in_table_cell0(39,
    "Three things no competitor on this list offers at our scale: 5,000+ certified OD trainers and consultants on demand, instructor-led training delivery mapped to every OD intervention across 100+ countries, and deep APAC, Middle East, and India-specific organizational development expertise — a market gap all 7 top competitors leave open.",
    "Three things no competitor on this list offers at our scale: 5,000+ certified organizational development consultants on demand, instructor-led training delivery mapped to every organizational development consulting intervention across 100+ countries, and deep APAC, Middle East, and India-specific organizational development consulting expertise — making Edstellar the top organizational development consulting provider. This is a market gap all 7 top competitors leave open.",
    ["certified organizational development consultants", "organizational development consulting", "top organizational development consulting provider"])
changes += 1

# ════════════════════════════════════════════════════════════
# T41: Gated Asset
# ════════════════════════════════════════════════════════════
print("\n=== T41: Gated Asset ===")

replace_in_table_cell0(41,
    "Not sure where to start your OD journey? Download Edstellar's free Organizational Development Readiness Checklist — a practical diagnostic tool that helps leadership teams assess structural gaps, culture alignment, and capability readiness before beginning a formal OD engagement.",
    "Not sure where to start your organizational development consulting journey? Download Edstellar's free Organizational Development Readiness Checklist — a practical diagnostic tool from our organizational development consulting experts that helps leadership teams assess structural gaps, culture alignment, and capability readiness before beginning a formal organizational development consultation.",
    ["organizational development consulting", "organizational development consulting experts", "organizational development consultation"])
changes += 1

# ════════════════════════════════════════════════════════════
# T43: CTA Banner
# ════════════════════════════════════════════════════════════
print("\n=== T43: CTA Banner ===")

replace_in_table_cell0(43,
    "Start with a diagnostic conversation. Edstellar's OD consultants will scope the right engagement model for your organization's complexity, urgency, and strategic priorities — no commitment required.",
    "Start with a diagnostic conversation. Edstellar's organizational development consultants will scope the right organizational development consulting engagement model for your organization's complexity, urgency, and strategic priorities — no commitment required. Outsource organizational development consulting to the experts.",
    ["organizational development consultants", "organizational development consulting", "outsource organizational development consulting"])
changes += 1

# ════════════════════════════════════════════════════════════
# T45: Contact Form section
# ════════════════════════════════════════════════════════════
print("\n=== T45: Contact Form ===")

replace_in_table_cell0(45,
    "Contact our OD consultants today to discuss how we can help you build a thriving, high-performance organization ready for future challenges.",
    "Contact our organizational development consultants today to discuss how our organizational development consulting services can help you build a thriving, high-performance organization ready for future challenges.",
    ["organizational development consultants", "organizational development consulting services"])
changes += 1

# ════════════════════════════════════════════════════════════
# T47: Resources
# ════════════════════════════════════════════════════════════
print("\n=== T47: Resources ===")

replace_in_table_cell0(47,
    "Deep dive into organizational development trends, expert insights, and practical frameworks.",
    "Deep dive into organizational development consulting trends, expert insights, and practical frameworks from Edstellar's organizational development consulting experts.",
    ["organizational development consulting", "organizational development consulting experts"])
changes += 1

# ════════════════════════════════════════════════════════════
# Paragraph 23: Body text about OD intervention
# ════════════════════════════════════════════════════════════
print("\n=== Paragraph 23 ===")

p23 = doc.paragraphs[23]
replace_in_paragraph_with_bold(p23,
    "Edstellar backs every OD intervention with 5,000+ certified trainers and 2,000+ instructor-led programs across 100+ locations, turning strategic recommendations into enterprise-wide execution.",
    "Edstellar backs every organizational development consulting intervention with 5,000+ certified organizational development consultants and 2,000+ instructor-led programs across 100+ locations, turning strategic recommendations into enterprise-wide execution.",
    ["organizational development consulting", "certified organizational development consultants"])
changes += 1

# ════════════════════════════════════════════════════════════
# Paragraph 25: Body text about OD initiatives
# ════════════════════════════════════════════════════════════
print("\n=== Paragraph 25 ===")

p25 = doc.paragraphs[25]
replace_in_paragraph_with_bold(p25,
    "Edstellar supports organizational development initiatives with global delivery capability, domain expertise, and structured execution across 100+ locations",
    "Edstellar's organizational development consulting services support organizational development initiatives with global delivery capability, domain expertise, and structured execution across 100+ locations",
    ["organizational development consulting services"])
changes += 1

# ════════════════════════════════════════════════════════════
# Now do a GLOBAL PASS: bold all keyword phrases wherever they exist
# This catches any keywords that already existed in the document
# ════════════════════════════════════════════════════════════
print("\n=== Global Bold Pass: Bolding all keyword phrases across entire document ===")

def bold_keywords_in_paragraph(paragraph, keywords):
    """Re-render paragraph with all keyword instances bolded."""
    full_text = paragraph.text
    if not full_text.strip():
        return False

    # Check if any keyword exists
    has_keyword = False
    for kw in keywords:
        if kw.lower() in full_text.lower():
            has_keyword = True
            break
    if not has_keyword:
        return False

    # Capture formatting
    font_name = font_size = font_color = font_bold = None
    if paragraph.runs:
        r = paragraph.runs[0]
        font_name = r.font.name
        font_size = r.font.size
        font_color = r.font.color.rgb if r.font.color and r.font.color.rgb else None
        font_bold = r.bold

    # Build pattern - longest first
    sorted_kws = sorted(keywords, key=len, reverse=True)
    pattern = '(' + '|'.join(re.escape(kw) for kw in sorted_kws) + ')'

    parts = re.split(pattern, full_text, flags=re.IGNORECASE)
    if len(parts) <= 1:
        return False

    # Clear runs
    p_el = paragraph._element
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for r in p_el.findall(f'{ns}r'):
        p_el.remove(r)

    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if font_color:
            run.font.color.rgb = font_color
        is_keyword = any(part.lower() == kw.lower() for kw in sorted_kws)
        if is_keyword:
            run.bold = True
        elif font_bold:
            run.bold = font_bold
    return True


# Global pass on all paragraphs
bold_count = 0
for i, para in enumerate(doc.paragraphs):
    if bold_keywords_in_paragraph(para, ALL_KEYWORDS):
        bold_count += 1
        print(f"  Bolded keywords in P{i}")

# Global pass on all table cells
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                if bold_keywords_in_paragraph(para, ALL_KEYWORDS):
                    bold_count += 1

print(f"\n  Total paragraphs/cells with keywords bolded in global pass: {bold_count}")

# ════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════
print(f"\n=== Saving document ===")
doc.save(OUTPUT_FILE)
print(f"Document saved to: {OUTPUT_FILE}")
print(f"Total targeted replacements: {changes}")
print("Done!")
