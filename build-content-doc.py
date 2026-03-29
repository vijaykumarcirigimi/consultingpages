#!/usr/bin/env python3
"""
Edstellar Developer Reference Document Builder

Parses a Design Scout report + keyword-optimized content file
and generates a rich Word document with every section tagged
by its library design module name.

Usage:
  python build-content-doc.py <report> <content> [--output <path>]

Example:
  python build-content-doc.py output/tna-report.md output/tna-optimized.md
  python build-content-doc.py output/tna-report.md output/tna-optimized.md -o output/tna-dev-ref.docx
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colors ──────────────────────────────────────────────
NAVY = RGBColor(0x2D, 0x2F, 0x6B)
LIME_TEXT = RGBColor(0x1A, 0x26, 0x00)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
BODY_TEXT = RGBColor(0x5C, 0x5E, 0x6E)
MUTED = RGBColor(0x8A, 0x8C, 0x9A)
RED = RGBColor(0xD9, 0x48, 0x48)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIME = RGBColor(0xC5, 0xE8, 0x26)
BLUE_LINK = RGBColor(0x7B, 0x93, 0xDB)


# ══════════════════════════════════════════════════════════════
#  PARSERS
# ══════════════════════════════════════════════════════════════

def parse_report(report_path):
    """Parse a Design Scout report and return page flow + metadata."""
    text = open(report_path, 'r', encoding='utf-8').read()

    # Page name
    m = re.search(r'\*\*Page:\*\*\s*(.+)', text)
    page_name = m.group(1).strip() if m else 'Consulting Page'

    # Date
    m = re.search(r'\*\*Date:\*\*\s*(.+)', text)
    date = m.group(1).strip() if m else ''

    # Section counts
    m = re.search(r'\*\*Sections identified:\*\*\s*(.+)', text)
    sections_info = m.group(1).strip() if m else ''

    m = re.search(r'\*\*Library matched:\*\*\s*(\d+)', text)
    lib_matched = m.group(1) if m else '0'

    m = re.search(r'\*\*Custom designs:\*\*\s*(\d+)', text)
    custom_count = m.group(1) if m else '0'

    # Page flow
    flow_match = re.search(r'## Recommended page flow\s*\n([\s\S]*?)(?=\n---|\n## |\Z)', text)
    page_flow = []
    if flow_match:
        for line in flow_match.group(1).strip().split('\n'):
            line = line.strip()
            if not line:
                continue

            # Library file: 1. `edstellar-xxx.html` — Section Name
            lib_m = re.match(r'^\d+\.\s*`([^`]+\.(?:html|md))`\s*(?:—|--|–|-)\s*(.+)', line)
            if lib_m:
                page_flow.append({
                    'type': 'library',
                    'file': lib_m.group(1),
                    'label': lib_m.group(2).strip()
                })
                continue

            # Custom: 7. `CUSTOM: #section-id` — Section Name
            cust_m = re.match(r'^\d+\.\s*`CUSTOM:\s*#([^`]+)`\s*(?:—|--|–|-)\s*(.+)', line)
            if cust_m:
                page_flow.append({
                    'type': 'custom',
                    'file': f'CUSTOM: #{cust_m.group(1)}',
                    'label': cust_m.group(2).strip()
                })
                continue

    # Notes
    notes_match = re.search(r'## Notes\s*\n([\s\S]*?)(?=\n---|\Z)', text)
    notes = []
    if notes_match:
        for line in notes_match.group(1).strip().split('\n'):
            line = line.strip()
            if line.startswith('- '):
                # Strip bold markers for clean text
                clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', line[2:])
                notes.append(clean)

    return {
        'page_name': page_name,
        'date': date,
        'sections_info': sections_info,
        'lib_matched': lib_matched,
        'custom_count': custom_count,
        'page_flow': page_flow,
        'notes': notes,
    }


def parse_content(content_path):
    """Parse a keyword-optimized content file into structured sections."""
    text = open(content_path, 'r', encoding='utf-8').read()

    # Extract SEO metadata block
    seo = {}
    seo_fields = [
        ('meta_title', r'\*\*Meta Title:\*\*\s*(.+?)(?:\n|$)'),
        ('meta_desc', r'\*\*Meta Description:\*\*\s*(.+?)(?:\n|$)'),
        ('h1', r'\*\*H1:\*\*\s*(.+?)(?:\n|$)'),
        ('canonical', r'\*\*Canonical URL:\*\*\s*(.+?)(?:\n|$)'),
        ('primary_kw', r'\*\*Primary Keywords?:\*\*\s*(.+?)(?:\n|$)'),
        ('secondary_kw', r'\*\*Secondary Keywords?:\*\*\s*(.+?)(?:\n|$)'),
    ]
    for key, pattern in seo_fields:
        m = re.search(pattern, text)
        if m:
            seo[key] = m.group(1).strip().rstrip('*').strip()

    # Split into sections by ## SECTION XX │ pattern
    section_splits = re.split(r'(?=## SECTION \d+)', text)
    sections = []

    for block in section_splits:
        block = block.strip()
        if not block:
            continue

        # Match section header
        header_m = re.match(r'## SECTION (\d+)\s*[│|]\s*(.+?)(?:\n|$)', block)
        if not header_m:
            continue

        sec_num = header_m.group(1).strip()
        sec_title = header_m.group(2).strip()

        # Get remaining content after header line
        body = block[header_m.end():].strip()

        # Remove the first --- separator if at the very end
        body = re.sub(r'\n---\s*$', '', body).strip()

        # Parse the body into structured fields
        fields = parse_section_body(body)

        sections.append({
            'num': sec_num,
            'title': sec_title,
            'body_raw': body,
            'fields': fields,
        })

    return {
        'seo': seo,
        'sections': sections,
    }


def parse_section_body(body):
    """Parse a section body into a list of (type, value) fields."""
    fields = []
    lines = body.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip SEO annotation lines
        if line.strip().startswith('🔍 SEO:') or line.strip().startswith('🔍'):
            i += 1
            continue

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # H3 heading (### ...)
        if line.strip().startswith('### '):
            heading_text = line.strip()[4:].strip()
            fields.append(('heading', heading_text))
            i += 1
            continue

        # Bold label with value on same line: **LABEL:** value
        bold_m = re.match(r'^\s*\*\*(.+?)(?::\*\*|\*\*:)\s*(.*)', line)
        if bold_m:
            label = bold_m.group(1).strip()
            value = bold_m.group(2).strip()

            # Check if next lines are bullet list under this label
            bullets = []
            j = i + 1
            while j < len(lines):
                bline = lines[j].strip()
                if bline.startswith('- ') or bline.startswith('* ') or bline.startswith('✓ ') or re.match(r'^- [✓✗●•]', bline):
                    bullet_text = re.sub(r'^[-*]\s*[✓✗●•]?\s*', '', bline).strip()
                    bullets.append(bullet_text)
                    j += 1
                elif bline == '' and j + 1 < len(lines) and (lines[j + 1].strip().startswith('- ') or lines[j + 1].strip().startswith('* ')):
                    j += 1
                else:
                    break

            if bullets:
                fields.append(('label_with_bullets', label, value, bullets))
                i = j
            elif value:
                fields.append(('label_value', label, value))
                i += 1
            else:
                # Value on next line(s) as paragraph
                para_lines = []
                j = i + 1
                while j < len(lines):
                    pline = lines[j].strip()
                    if not pline or pline.startswith('**') or pline.startswith('### ') or pline.startswith('## ') or pline.startswith('🔍'):
                        break
                    para_lines.append(pline)
                    j += 1
                if para_lines:
                    fields.append(('label_value', label, ' '.join(para_lines)))
                else:
                    fields.append(('label_value', label, ''))
                i = j
            continue

        # Bullet list (standalone)
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            bullets = []
            while i < len(lines):
                bline = lines[i].strip()
                if bline.startswith('- ') or bline.startswith('* '):
                    bullet_text = re.sub(r'^[-*]\s*[✓✗●•]?\s*', '', bline).strip()
                    bullets.append(bullet_text)
                    i += 1
                else:
                    break
            fields.append(('bullets', bullets))
            continue

        # Italic line (*text*)
        italic_m = re.match(r'^\s*\*([^*]+)\*\s*$', line)
        if italic_m:
            fields.append(('italic', italic_m.group(1).strip()))
            i += 1
            continue

        # Plain paragraph text
        para_lines = [line.strip()]
        j = i + 1
        while j < len(lines):
            pline = lines[j].strip()
            if not pline or pline.startswith('**') or pline.startswith('### ') or pline.startswith('## ') or pline.startswith('- ') or pline.startswith('* ') or pline.startswith('🔍'):
                break
            para_lines.append(pline)
            j += 1
        text = ' '.join(para_lines)
        if text:
            fields.append(('paragraph', text))
        i = j

    return fields


# ══════════════════════════════════════════════════════════════
#  DOCUMENT BUILDER HELPERS
# ══════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_tag(doc, index, section_label, design_file, is_custom=False):
    """Add the navy design module tag bar for a section."""
    # Divider
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("_" * 80)
    run.font.color.rgb = RGBColor(0xDC, 0xDE, 0xE5)
    run.font.size = Pt(8)

    # Section header
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"S{index:02d}  |  ")
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    run.font.bold = True
    run = p.add_run(section_label)
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    run.font.bold = True

    # Design module tag
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cell = table.cell(0, 0)
    set_cell_shading(cell, "2D2F6B")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    label_text = "  CUSTOM DESIGN  " if is_custom else "  DESIGN MODULE  "
    run = p.add_run(label_text)
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = LIME if not is_custom else RGBColor(0xF0, 0xA8, 0x62)

    cell = table.cell(0, 1)
    set_cell_shading(cell, "2D2F6B")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {design_file}  ")
    run.font.size = Pt(9)
    run.font.color.rgb = WHITE
    run.font.bold = True

    # Spacer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_field_label(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f" {label} ")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.all_caps = True


def add_text(doc, text, bold=False, italic=False, color=None, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    # Handle inline bold markers **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.size = Pt(10)
            run.font.color.rgb = color or DARK_TEXT
            run.font.bold = True
            run.font.italic = italic
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = color or BODY_TEXT
            run.font.bold = bold
            run.font.italic = italic
    return p


def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(1.5)
    p.clear()
    # Handle inline bold
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT
            run.font.bold = True
        else:
            run = p.add_run(part)
            run.font.size = Pt(10)
            run.font.color.rgb = color or BODY_TEXT


def render_section_fields(doc, fields):
    """Render parsed fields into the Word document."""
    for field in fields:
        ftype = field[0]

        if ftype == 'heading':
            add_field_label(doc, field[1])

        elif ftype == 'label_value':
            label, value = field[1], field[2]
            # Categorize by known label patterns
            label_upper = label.upper()

            if any(k in label_upper for k in ['IMAGE PLACEHOLDER', 'IMAGE:', 'IMAGE NOTE']):
                # Image placeholder
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run("[IMAGE] ")
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = BLUE_LINK
                run = p.add_run(value)
                run.font.size = Pt(9)
                run.font.color.rgb = MUTED
                run.font.italic = True

            elif any(k in label_upper for k in ['PRIMARY CTA', 'SECONDARY CTA', 'CTA COPY', 'CTA:', 'CTA BUTTON', 'BOTTOM CTA', 'SUBMIT']):
                # CTA button
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(f"[{label}] ")
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = LIME_TEXT
                run = p.add_run(value)
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = DARK_TEXT

            elif any(k in label_upper for k in ['STAT ', 'STAT:', 'METRIC']):
                # Stat value
                add_field_label(doc, label)
                add_text(doc, value, bold=True, color=NAVY)

            elif any(k in label_upper for k in ['SOURCE', 'NOTE:', 'FORMAT:', 'PRIORITY:', 'FUNNEL']):
                # Metadata note
                add_text(doc, f"{label}: {value}", italic=True, color=MUTED)

            else:
                # Generic label: value
                add_field_label(doc, label)
                if value:
                    add_text(doc, value)

        elif ftype == 'label_with_bullets':
            label, value, bullets = field[1], field[2], field[3]
            add_field_label(doc, label)
            if value:
                add_text(doc, value)
            for b in bullets:
                add_bullet(doc, b)

        elif ftype == 'bullets':
            for b in field[1]:
                add_bullet(doc, b)

        elif ftype == 'paragraph':
            add_text(doc, field[1])

        elif ftype == 'italic':
            add_text(doc, field[1], italic=True, color=MUTED)


# ══════════════════════════════════════════════════════════════
#  MAIN DOCUMENT GENERATOR
# ══════════════════════════════════════════════════════════════

def build_document(report_path, content_path, output_path):
    report = parse_report(report_path)
    content = parse_content(content_path)

    page_name = report['page_name']
    page_flow = report['page_flow']
    seo = content['seo']
    sections = content['sections']

    doc = Document()

    # ── Page setup ──
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = BODY_TEXT

    # ══════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("EDSTELLAR")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    run.font.letter_spacing = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    run = p.add_run(page_name)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Consulting Services")
    run.font.size = Pt(24)
    run.font.color.rgb = LIME_TEXT
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    run = p.add_run("Developer Content Reference Document")
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("Each section tagged with its design module from the Edstellar HTML library.\nUse this document to build the page manually.")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    if report['date']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(24)
        run = p.add_run(f"{report['date']}  |  {report['sections_info']}")
        run.font.size = Pt(10)
        run.font.color.rgb = NAVY
        run.font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # SEO METADATA PAGE
    # ══════════════════════════════════════════════════════
    if seo:
        p = doc.add_paragraph()
        run = p.add_run("SEO METADATA")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = NAVY

        seo_rows = [
            ("Meta Title", seo.get('meta_title', '')),
            ("Meta Description", seo.get('meta_desc', '')),
            ("H1", seo.get('h1', '')),
            ("Canonical URL", seo.get('canonical', '')),
            ("Primary Keywords", seo.get('primary_kw', '')),
            ("Secondary Keywords", seo.get('secondary_kw', '')),
        ]
        seo_rows = [(k, v) for k, v in seo_rows if v]

        if seo_rows:
            table = doc.add_table(rows=len(seo_rows) + 1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT

            for ci, hdr in enumerate(["Field", "Value"]):
                cell = table.cell(0, ci)
                set_cell_shading(cell, "2D2F6B")
                p = cell.paragraphs[0]
                run = p.add_run(hdr)
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = WHITE

            for ri, (field, value) in enumerate(seo_rows, 1):
                cell = table.cell(ri, 0)
                p = cell.paragraphs[0]
                run = p.add_run(field)
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = NAVY

                cell = table.cell(ri, 1)
                p = cell.paragraphs[0]
                run = p.add_run(value)
                run.font.size = Pt(9)
                run.font.color.rgb = BODY_TEXT

        doc.add_paragraph()

    # ══════════════════════════════════════════════════════
    # PAGE FLOW OVERVIEW TABLE
    # ══════════════════════════════════════════════════════
    p = doc.add_paragraph()
    run = p.add_run("PAGE FLOW OVERVIEW")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    run = p.add_run(f"{len(page_flow)} sections in order. Each row shows the section name and the HTML design module file to use from the library/ folder.")
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    table = doc.add_table(rows=len(page_flow) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, hdr in enumerate(["#", "Section", "Design Module / Custom"]):
        cell = table.cell(0, ci)
        set_cell_shading(cell, "2D2F6B")
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = WHITE

    for ri, item in enumerate(page_flow, 1):
        vals = [f"{ri:02d}", item['label'], item['file']]
        for ci, val in enumerate(vals):
            cell = table.cell(ri, ci)
            if ri % 2 == 0:
                set_cell_shading(cell, "F4F5F7")
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.color.rgb = DARK_TEXT if ci < 2 else NAVY
            run.font.bold = ci == 2

    # Notes from report
    if report['notes']:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("NOTES")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY

        for note in report['notes']:
            add_bullet(doc, note, color=MUTED)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # SECTION-BY-SECTION CONTENT
    # ══════════════════════════════════════════════════════

    # Build a mapping from content section number to page flow entry
    # Content sections are numbered 01, 02, ... and page flow entries
    # include the same sections plus structural ones (SEO metadata, footer)
    # We match by index: content section i maps to page flow entry
    # that has the same position (skipping seo-metadata-block-template.md)

    # Filter page flow to get visual sections (skip seo-metadata template)
    visual_flow = [item for item in page_flow if not item['file'].endswith('.md')]

    for i, section in enumerate(sections):
        sec_num = int(section['num'])

        # Find matching page flow entry
        flow_item = None
        # Try to match by index (content section 01 -> first visual flow entry, etc.)
        flow_idx = sec_num - 1
        if flow_idx < len(visual_flow):
            flow_item = visual_flow[flow_idx]

        if flow_item:
            is_custom = flow_item['type'] == 'custom'
            add_section_tag(doc, sec_num, section['title'], flow_item['file'], is_custom)
        else:
            # No flow match, just show section header
            add_section_tag(doc, sec_num, section['title'], '(not mapped)', is_custom=False)

        # Render the section content
        render_section_fields(doc, section['fields'])

    # ── Save ──
    doc.save(output_path)
    print(f"\n  Document saved: {output_path}")
    print(f"  Page: {page_name}")
    print(f"  Sections: {len(sections)} content | {len(page_flow)} flow entries")
    print(f"  Library matched: {report['lib_matched']} | Custom: {report['custom_count']}\n")


# ══════════════════════════════════════════════════════════════
#  CLI
# ══════════════════════════════════════════════════════════════

def main():
    args = sys.argv[1:]

    if len(args) < 2 or '--help' in args or '-h' in args:
        print("""
  Edstellar Developer Reference Document Builder
  ================================================

  Generates a rich Word document from a Design Scout report
  + keyword-optimized content file, with each section tagged
  by its library design module.

  Usage:
    python build-content-doc.py <report> <content> [--output <path>]

  Options:
    --output, -o   Output file path (default: output/<slug>-developer-reference.docx)
    --help, -h     Show this help

  Example:
    python build-content-doc.py output/tna-report.md output/tna-optimized.md
""")
        sys.exit(0)

    report_path = os.path.abspath(args[0])
    content_path = os.path.abspath(args[1])

    if not os.path.exists(report_path):
        print(f"  Error: Report file not found: {report_path}")
        sys.exit(1)
    if not os.path.exists(content_path):
        print(f"  Error: Content file not found: {content_path}")
        sys.exit(1)

    # Determine output path
    output_path = None
    out_idx = -1
    for flag in ['--output', '-o']:
        if flag in args:
            out_idx = args.index(flag)
            break
    if out_idx != -1 and out_idx + 1 < len(args):
        output_path = os.path.abspath(args[out_idx + 1])
    else:
        # Derive from report filename
        base = os.path.basename(report_path)
        slug = base.replace('-report.md', '').replace('.md', '')
        output_path = os.path.join(os.path.dirname(report_path), f"{slug}-developer-reference.docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    print("\n  Edstellar Developer Reference Document Builder")
    print("  ================================================")
    build_document(report_path, content_path, output_path)
    print("  Done.\n")


if __name__ == "__main__":
    main()
