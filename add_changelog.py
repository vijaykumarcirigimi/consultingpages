from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("C:/Users/Edstellar/Documents/Consulting Page Builder/content/Final Pages/Talent_Assessments_Final.docx")

# Add page break
doc.add_page_break()

# Add heading (manually since Heading 1 style doesn't exist in this doc)
heading_para = doc.add_paragraph()
heading_run = heading_para.add_run("Keyword Optimization Change Log")
heading_run.bold = True
heading_run.font.size = Pt(18)
heading_run.font.color.rgb = RGBColor(0x2D, 0x2F, 0x6B)  # Navy brand color

# Add intro paragraph
intro = doc.add_paragraph(
    "Summary of keyword changes made to this document. "
    "All added keywords are bolded in the content above."
)

# Define the change rows grouped by section area
rows_data = [
    # Group: Meta Titles/Descriptions/H1s (T5-T7)
    ("T5-T7: Meta Titles & Descriptions",
     "Meta descriptions without 'talent assessment provider' keyword",
     "Wove 'talent assessment provider' into meta descriptions"),
    ("T5-T7: Meta Titles & Descriptions",
     "Meta descriptions without 'talent assessment solutions' keyword",
     "Wove 'talent assessment solutions' into meta descriptions"),
    ("T5-T7: H1 Variations",
     "H1 variations without 'talent assessment consulting' keyword",
     "Wove 'talent assessment consulting' into H1 variations"),

    # Group: Supporting Keywords (T9-T13)
    ("T9: Supporting Keywords Table 1",
     "Table had fewer keyword rows",
     "Added 2 rows: 'Talent Assessment Consulting', 'Custom Talent Assessment Services'"),
    ("T10: Supporting Keywords Table 2",
     "Table had fewer keyword rows",
     "Added 2 rows: 'Talent Assessment Service Provider', 'Hire Talent Assessment Consultant'"),
    ("T11: Supporting Keywords Table 3",
     "Table had fewer keyword rows",
     "Added 2 rows: 'Talent Assessment Partner', 'Professional Talent Assessment Services'"),
    ("T12: Supporting Keywords Table 4",
     "Table had fewer keyword rows",
     "Added 1 row: '360-Degree Feedback and Talent Assessment Services'"),
    ("T13: Supporting Keywords Table 5",
     "Table had fewer keyword rows",
     "Added 7 rows: 'Enterprise-Grade Talent Assessment', 'Managed Talent Assessment Services', 'Outsource Talent Assessment', 'Talent Assessment Firm', 'Talent Assessment Vendor', 'Talent Assessment Agency', 'Talent Assessment Experts'"),

    # Group: Service Descriptions (T15-T21)
    ("T15: Service Description 1",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),
    ("T16: Service Description 2",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),
    ("T17: Service Description 3",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),
    ("T18: Service Description 4",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),
    ("T19: Service Description 5",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),
    ("T20: Service Description 6",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),
    ("T21: Service Description 7",
     "Source line without talent assessment qualifier",
     "Injected 'talent assessment provider/company/solutions/consultants' into service description source line"),

    # Group: Page Structure (T23)
    ("T23: Page Structure",
     "Section descriptions without 'talent assessment' qualifier (e.g., '7 service offering cards')",
     "Added 'talent assessment' qualifiers to section descriptions (e.g., '7 talent assessment service offering cards')"),

    # Group: Page Headings (T25)
    ("T25: Page Heading 1",
     "Heading without 'talent assessment' keyword",
     "Updated to: 'How Our 5-Step Talent Assessment Process Works'"),
    ("T25: Page Heading 2",
     "Heading without 'talent assessment' keyword",
     "Updated to: 'What Sets Edstellar Apart as a Talent Assessment Partner'"),
    ("T25: Page Heading 3",
     "Heading without 'talent assessment' keyword",
     "Updated heading to include 'talent assessment'"),
    ("T25: Page Heading 4",
     "Heading without 'talent assessment' keyword",
     "Updated heading to include 'talent assessment'"),
    ("T25: Page Heading 5",
     "Heading without 'talent assessment' keyword",
     "Updated heading to include 'talent assessment'"),
    ("T25: Page Heading 6",
     "Heading without 'talent assessment' keyword",
     "Updated heading to include 'talent assessment'"),
    ("T25: Page Heading 7",
     "Heading without 'talent assessment' keyword",
     "Updated heading to include 'talent assessment'"),

    # Group: Content Guidance (T27-T36)
    ("T27: Content Guidance Section 1",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T28: Content Guidance Section 2",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T29: Content Guidance Section 3",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T30: Content Guidance Section 4",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T31: Content Guidance Section 5",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T32: Content Guidance Section 6",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T33: Content Guidance Section 7",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T34: Content Guidance Section 8",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T35: Content Guidance Section 9",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),
    ("T36: Content Guidance Section 10",
     "Section keyword list without talent assessment targets",
     "Added keyword targets to section keyword list"),

    # Group: CTAs (T38-T41)
    ("T38: CTA 1",
     "CTA without 'talent assessment' keyword",
     "Updated to: 'Start Your Free Talent Assessment Consultation'"),
    ("T39: CTA 2",
     "CTA without 'talent assessment' keyword",
     "Updated to: 'Download Our Talent Assessment Tools Guide'"),
    ("T40: CTA 3",
     "CTA without 'talent assessment' keyword",
     "Updated CTA with 'talent assessment' keyword"),
    ("T41: CTA 4",
     "CTA without 'talent assessment' keyword",
     "Updated CTA with 'talent assessment' keyword"),

    # Group: Assets (T43-T46)
    ("T43: Asset Description 1",
     "Asset description without 'talent assessment' keyword",
     "Wove 'talent assessment' keywords into asset description"),
    ("T44: Asset Description 2",
     "Asset description without 'talent assessment' keyword",
     "Wove 'talent assessment' keywords into asset description"),
    ("T45: Asset Description 3",
     "Asset description without 'talent assessment' keyword",
     "Wove 'talent assessment' keywords into asset description"),
    ("T46: Asset Description 4",
     "Asset description without 'talent assessment' keyword",
     "Wove 'talent assessment' keywords into asset description"),

    # Group: Quick Reference (T48-T50)
    ("T48-T50: Quick Reference Tables",
     "Keyword instances not bolded in summary tables",
     "Bolded all keyword instances in summary tables"),
]

# Create table
table = doc.add_table(rows=1 + len(rows_data), cols=3)
# Apply borders manually since Table Grid style doesn't exist
from docx.oxml import OxmlElement
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)
set_table_borders(table)

# Set column widths
for row in table.rows:
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(2.6)
    row.cells[2].width = Inches(2.6)

# Header row
headers = ["Section", "Earlier Content", "Updated Content"]
for i, header_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(header_text)
    run.bold = True
    run.font.size = Pt(9)
    # Set header cell shading to light gray
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): "D9E2F3"
    })
    shading.append(shading_elem)

# Data rows
for row_idx, (section, earlier, updated) in enumerate(rows_data, start=1):
    row = table.rows[row_idx]
    for col_idx, text in enumerate([section, earlier, updated]):
        cell = row.cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(8)
        if col_idx == 0:
            run.bold = True

# Save
doc.save("C:/Users/Edstellar/Documents/Consulting Page Builder/content/Final Pages/Talent_Assessments_Final.docx")
print(f"Done. Added {len(rows_data)} rows to the change log table.")
